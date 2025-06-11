from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, black
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas
from reportlab.pdfbase.pdfmetrics import stringWidth
import json
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from app.constants.colors import COLORS
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def generate_pdf_summary(json_string, file_buffer):
    data = json.loads(json_string)
    c = canvas.Canvas(file_buffer, pagesize=A4)
    width, height = A4
    margin = 50
    usable_width = width - 2 * margin
    y = height - margin
    line_height = 14

    colors = {i + 1: HexColor(color_def['hexa']) for i, color_def in enumerate(COLORS)}

    logo_path = "app/uistatic/logos/scanlitt_bleu_bloc.png"
    logo_width = 101
    logo_height = 56
    logo_x = (width - logo_width) / 2
    logo_y = y - logo_height

    c.drawImage(ImageReader(logo_path), logo_x, logo_y, width=logo_width, height=logo_height, mask='auto')
    y = logo_y - 20 

    def wrap_text(text, font_name, font_size, max_width):
        words = text.split()
        lines = []
        current_line = ""
        for word in words:
            test_line = f"{current_line} {word}".strip()
            if stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                lines.append(current_line)
                current_line = word
        if current_line:
            lines.append(current_line)
        return lines

    def draw_line(text, color=black, size=12, bold=False, spacer=0):
        nonlocal y
        font_name = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font_name, size)
        c.setFillColor(color)
        wrapped_lines = wrap_text(text, font_name, size, usable_width)
        for line in wrapped_lines:
            if y < margin:
                c.showPage()
                y = height - margin
                c.setFont(font_name, size)
                c.setFillColor(color)
            c.drawString(margin, y, line)
            y -= line_height
        y -= spacer

    # Header
    draw_line(data["main_title"], size=16, bold=True, spacer=20)

    for cluster in data.get("clusters", []):
        color = colors.get(cluster["cluster"], black)
        draw_line(f"Cluster {cluster['cluster']} — {cluster['common_theme']}", color=color, size=13, bold=True)
        draw_line(cluster["summary"], spacer=10)
        draw_line("Subthemes", bold=True)

        for sub in cluster["subthemes"]:
            citations = ", ".join(str(cit) for cit in sub["subtheme_citations"])
            draw_line(f"{sub['subtheme_type']}: {citations}", spacer=10)

        if cluster.get("outliers"):
            draw_line("Citations that do not belong to any subtheme", bold=True)
            for outlier in cluster["outliers"]:
                draw_line(str(outlier), spacer=10)

    if "cluster_zero_ref" in data:
        draw_line("Cluster Zero Reference", size=13, bold=True, spacer=10)
        for cluster_ref in data["cluster_zero_ref"]:
            fitted = cluster_ref["fitted_cluster"]
            citations = ", ".join(str(c) for c in cluster_ref["citations_zero"])
            draw_line(f"Fitted Cluster: {fitted} — Citations Zero: {citations}", spacer=10)
    c.save()
    return file_buffer


def get_color_rgb(cluster_index):
    color = COLORS[cluster_index - 1]  # index starts at 0
    return RGBColor(color['red'], color['green'], color['blue'])

def generate_docx_summary(json_string, filename):
    data = json.loads(json_string)
    doc = Document()

    logo_path = "app/uistatic/logos/scanlitt_bleu_bloc.png"
    logo_width_cm = 3

    logo_paragraph = doc.add_paragraph()
    run = logo_paragraph.add_run()
    run.add_picture(logo_path, width=Cm(logo_width_cm))
    logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading(data["main_title"], level=1)
    doc.add_paragraph(f"Modèle : {data.get('result_model', 'N/A')}")
    doc.add_paragraph("")

    # Clusters normaux
    for cluster in data.get("clusters", []):
        color = get_color_rgb(cluster["cluster"])
        p = doc.add_paragraph()
        run = p.add_run(f"Cluster {cluster['cluster']} — {cluster['common_theme']}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = color

        doc.add_paragraph(cluster["summary"])
        doc.add_paragraph("Subthemes", style='Heading 3')

        for sub in cluster["subthemes"]:
            citations = ", ".join(str(cit) for cit in sub["subtheme_citations"])
            doc.add_paragraph(f"{sub['subtheme_type']}: {citations}")

        if cluster.get("outliers"):
            doc.add_paragraph("Citations that do not belong to any subtheme", style='Heading 3')
            for outlier in cluster["outliers"]:
                doc.add_paragraph(str(outlier))

        doc.add_paragraph("")

    if "cluster_zero_ref" in data:
        doc.add_heading("Cluster Zero Reference", level=2)
        for cluster_ref in data["cluster_zero_ref"]:
            fitted = cluster_ref["fitted_cluster"]
            citations = ", ".join(str(c) for c in cluster_ref["citations_zero"])
            doc.add_paragraph(f"Fitted Cluster: {fitted} — Citations Zero: {citations}")

    # Sauvegarde
    doc.save(filename)
    return filename
