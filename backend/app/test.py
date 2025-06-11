from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor, black
from reportlab.pdfgen import canvas
from reportlab.pdfbase.pdfmetrics import stringWidth
import json
from docx import Document
from docx.shared import Pt, RGBColor
from constants.colors import COLORS

filename_pdf="application/app/app/technostress_summary.pdf"
filename_docx="application/app/app/technostress_summary.docx"
with open("application/app/app/test.json") as f:
    data = json.load(f)

def generate_pdf_summary(data, filename):
    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4
    margin = 50
    usable_width = width - 2 * margin
    y = height - margin
    line_height = 14

    colors = {i + 1: HexColor(color_def['hexa']) for i, color_def in enumerate(COLORS)}
    print('ici ça marche')
    def wrap_text(text, font_name, font_size, max_width):
        words = text.split()
        lines = []
        current_line = ""

        for word in words:
            test_line = f"{current_line} {word}".strip()
            if stringWidth(test_line, font_name, font_size) <= max_width:
                current_line = test_line
            else:
                lines.append(current_line)
                current_line = word
        if current_line:
            lines.append(current_line)
        return lines

    def draw_line(text, color=black, size=12, bold=False, spacer=0):
        nonlocal y
        font_name = "Helvetica-Bold" if bold else "Helvetica"
        c.setFont(font_name, size)
        c.setFillColor(color)

        wrapped_lines = wrap_text(text, font_name, size, usable_width)
        for line in wrapped_lines:
            if y < margin:
                c.showPage()
                y = height - margin
                c.setFont(font_name, size)
                c.setFillColor(color)
            c.drawString(margin, y, line)
            y -= line_height

        y -= spacer

    # Header
    draw_line(data["main_title"], size=16, bold=True, spacer=20)
    #draw_line(data["sub_title"], size=14, bold=True)


    # Clusters
    for cluster in data["clusters"]:
        color = colors.get(cluster["cluster"], black)
        draw_line(f"Cluster {cluster['cluster']} — {cluster['common_theme']}", color=color, size=13, bold=True)
        draw_line(cluster["summary"], spacer=10)
        draw_line("Subthemes", bold=True)

        for sub in cluster["subthemes"]:
            citations = ", ".join(sub["subtheme_citations"])
            draw_line(f"{sub['subtheme_type']}: {citations}", spacer=10)

        if cluster.get("outliers"):
            draw_line("Citations that do not belong to any subtheme", bold=True)
            for outlier in cluster["outliers"]:
                draw_line(outlier, spacer=10)

    draw_line(f"Modèle : {data['result_model']}")
    c.save()
    return filename


def get_color_rgb(cluster_index):
    color = COLORS[cluster_index - 1]  # index starts at 0
    return RGBColor(color['red'], color['green'], color['blue'])

def generate_docx_summary(data, filename):
    doc = Document()

    # Titre principal
    title = doc.add_heading(data["main_title"], level=1)
    title.alignment = 0

    doc.add_paragraph(f"Modèle : {data['result_model']}")
    #doc.add_paragraph(f"Coût : {data['couts']}")
    #doc.add_paragraph(f"Prompt tokens : {data['prompt_token']}")
    #doc.add_paragraph(f"Completion tokens : {data['completion_token']}")
    #doc.add_paragraph(f"Temps de réponse API : {data.get('api_response_time', 'N/A')} s")
    doc.add_paragraph("")  # Espace

    # Clusters
    for cluster in data["clusters"]:
        color = get_color_rgb(cluster["cluster"])

        # Cluster title
        p = doc.add_paragraph()
        run = p.add_run(f"Cluster {cluster['cluster']} — {cluster['common_theme']}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = color

        # Summary
        doc.add_paragraph(cluster["summary"])

        # Subthemes
        doc.add_paragraph("Subthemes", style='Heading 3')
        for sub in cluster["subthemes"]:
            citations = ", ".join(sub["subtheme_citations"])
            doc.add_paragraph(f"{sub['subtheme_type']}: {citations}")

        # Outliers
        if cluster.get("outliers"):
            doc.add_paragraph("Citations that do not belong to any subtheme", style='Heading 3')
            for outlier in cluster["outliers"]:
                doc.add_paragraph(outlier)

        doc.add_paragraph("")  # Espace entre clusters

    # Sauvegarde
    doc.save(filename)
    return filename

generate_pdf_summary(data, filename_pdf)
generate_docx_summary(data, filename_docx)